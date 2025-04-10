import os
import docx
from pdf2docx import Converter
from ner_model import ner_model


def convert_pdf_to_docx(pdf_file_path, docx_file_path):
    try:
        cv = Converter(pdf_file_path)
        cv.convert(docx_file_path, start=0, end=None)
        cv.close()
        print(f"DOCX saved to: {docx_file_path}")
    except Exception as e:
        print(f"Conversion failed: {e}")

def word_parser(file_path):
    doc = docx.Document(file_path)

    lines = []

    paragraph_count = len([p for p in doc.paragraphs if p.text.strip()])
    table_count = len(doc.tables)

    # Paragraphs
    # if paragraph_count:
    #     print("Parsing from paragraph.")
    #     for para in doc.paragraphs:
    #         print(para.text)
    #         if para.text.strip():
    #             lines.append(para.text.strip())
    
    # Tables
    if table_count:
        print("Parsing from table.\n")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        lines.append(cell_text)
    return lines

def find_skill_section(lines):
    max_lines = 15
    skills_text = []
    capture_skills_section =False
    for line in lines:
        if "skills" in line.lower():
            capture_skills_section = True

        if capture_skills_section:
            skills_text.append(line)
        
        if len(skills_text) > max_lines:
            break

    text = "\n".join(skills_text)
    return text

def extract_skills_from_word(file):
    file_path = f"data/sample_resumes/{file}"

    ext = os.path.splitext(file_path)[1].lower()
    if ext != ".pdf" and ext != ".docx":
        raise ValueError("Unsupported file type")

    if ext == ".pdf":
        docx_file_path = file_path.split(".")[0]+".docx"
        convert_pdf_to_docx(file_path, docx_file_path)
        file_path = docx_file_path
    
    data = word_parser(file_path)
    text = find_skill_section(data)
    ner_model(text)




if __name__ == "__main__":
    # file1 = "data-engineer-resume.docx"
    # file2 = "data-scientist-resume-example.pdf"
    file = "data-scientist-resume-example.docx"

    extract_skills_from_word(file)







    



    




 


