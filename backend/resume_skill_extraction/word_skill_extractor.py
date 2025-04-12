import docx
from transformers import pipeline
from difflib import get_close_matches


SKILL_KEYWORDS = [
    "skills", "technical skills", "core competencies",
    "technologies", "tools", "expertise", "proficiencies",
    "it skills", "additional skills"
]

file = "data-engineer-resume.docx"
file_path = f"data/sample_resumes/{file}"

known_skills = ["python", "sql", "excel", "tableau", "data analysis", "aws", "pandas", "keras"]


def get_text():
    lines = word_parser(file_path)
    return "\n".join(lines)

def get_skills(file_path, known_skills):
    # Step 1: Parse lines from Word document
    data = word_parser(file_path)

    # Step 2: Extract the likely skill section (optional)
    skill_section_text = enhanced_skill_section(data)

    # Step 3: Use NER if you have a model
    entities = ner_model(skill_section_text)  # Optional step, depends on your model

    # Step 4: Match close skills from known list
    matched_skills = match_close_skills(data, known_skills)

    return matched_skills


def word_parser(file_path):
    doc = docx.Document(file_path)

    lines = []

    paragraph_count = len([p for p in doc.paragraphs if p.text.strip()])
    table_count = len(doc.tables)

    # Paragraphs
    if paragraph_count:
        print("Parsing from paragraph.")
        for para in doc.paragraphs:
            print(para.text)
            if para.text.strip():
                lines.append(para.text.strip())
    
    # Tables
    if table_count:
        print("Parsing from table.\n")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        lines.append(cell_text)
    return lines

def find_skill_section(lines):
    max_lines = 20
    skills_text = []
    capture_skills_section =False
    for line in lines:
        if line.lower() in SKILL_KEYWORDS:
            capture_skills_section = True

        if capture_skills_section:
            skills_text.append(line)
        
        if len(skills_text) > max_lines:
            break

    text = "\n".join(skills_text)
    # print(text)

    return text

def ner_model(text):
    ner = pipeline("ner", model="jjzha/jobbert_knowledge_extraction", grouped_entities=True)

    print(text)
    results = ner(text)
    print(results)


    skills = reconstruct_bio_entities(results)

    print(" Extracted Skills/Knowledge:")
    for skill in skills:
        print("-", skill)


def reconstruct_bio_entities(entities):
    skills = []
    current_phrase = []

    for ent in entities:
        word = ent["word"]
        label = ent["entity_group"]

        # Handle subword tokens (e.g., ##s)
        if word.startswith("##"):
            word = word[2:]
            if current_phrase:
                current_phrase[-1] += word
            continue

        if label == "B":
            if current_phrase:
                skills.append(" ".join(current_phrase))
            current_phrase = [word]
        elif label == "I":
            current_phrase.append(word)
            
    # Append the last phrase
    if current_phrase:
        skills.append(" ".join(current_phrase))

    return list(set(skills))  # Remove duplicates
    


def match_close_skills(text_lines, known_skills):
    matched = set()
    for line in text_lines:
        words = re.findall(r'\b\w[\w\-]+\b', line.lower())
        for word in words:
            matches = get_close_matches(word, known_skills, n=1, cutoff=0.85)
            if matches:
                matched.add(matches[0])
    return list(matched)

def enhanced_skill_section(lines, keywords=SKILL_KEYWORDS):
    text = "\n".join(lines).lower()
    for keyword in keywords:
        if keyword in text:
            idx = text.find(keyword)
            return text[idx:idx+500]  # extract section around keyword
    return "" 


if __name__ == "__main__":
    get_skills()



    




 


